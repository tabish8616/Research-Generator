import streamlit as st
import json
import re
import io
import zipfile
from datetime import datetime

import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
)

try:
    import google.generativeai as genai
    GENAI_AVAILABLE = True
except Exception:
    GENAI_AVAILABLE = False

MODEL_NAME = "gemini-2.5-flash-lite"

CITATION_STYLES = ["APA", "MLA", "Chicago", "Harvard", "IEEE"]
RESEARCH_TYPES = ["Quantitative", "Qualitative", "Mixed Methods"]
RESEARCH_LEVELS = ["Undergraduate", "Postgraduate", "PhD"]
JOURNAL_FORMATS = ["Generic Academic", "IEEE", "Springer", "Elsevier", "College Project"]
METHOD_APPROACHES = ["Quantitative", "Qualitative", "Mixed"]
METHOD_DESIGNS = ["Descriptive", "Exploratory", "Experimental"]

STEP_TITLES = [
    "1. Topic", "2. Problem", "3. Objectives", "4. Literature",
    "5. Gap", "6. Methodology", "7. Questionnaire", "8. Results",
    "9. Conclusion", "10. References", "Final Paper", "Quality Check",
]


def inject_css():
    st.markdown(
        """
        <style>
        .main { background-color: #F5F7FA; }
        h1, h2, h3 { color: #0D3B66; font-family: 'Georgia', serif; }
        .app-header {
            padding: 1.4rem 1.8rem;
            background: linear-gradient(135deg, #0D3B66 0%, #1565C0 100%);
            border-radius: 14px;
            color: white;
            margin-bottom: 1.2rem;
            box-shadow: 0 4px 14px rgba(13, 59, 102, 0.25);
        }
        .app-header h1 { color: white; margin: 0; font-size: 1.7rem; }
        .app-header p { color: #DCE8F7; margin: 0.3rem 0 0 0; font-size: 0.95rem; }
        .stTabs [data-baseweb="tab-list"] { gap: 4px; }
        .stTabs [data-baseweb="tab"] {
            background-color: #EAF1FB;
            border-radius: 8px 8px 0 0;
            padding: 8px 14px;
            font-weight: 500;
            color: #0D3B66;
        }
        .stTabs [aria-selected="true"] {
            background-color: #1565C0 !important;
            color: white !important;
        }
        div[data-testid="stMetricValue"] { color: #0D3B66; }
        .verify-box {
            background-color: #FFF4E5;
            border-left: 5px solid #E8A33D;
            padding: 0.7rem 1rem;
            border-radius: 8px;
            font-size: 0.88rem;
            color: #6B4A12;
            margin: 0.6rem 0;
        }
        .privacy-box {
            background-color: #E9F3FF;
            border-left: 4px solid #1565C0;
            padding: 0.6rem 0.9rem;
            border-radius: 8px;
            font-size: 0.8rem;
            color: #0D3B66;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def init_session_state():
    defaults = {
        "api_key": "",
        "citation_style": "APA",
        "word_count_target": 8000,
        "page_count_target": 25,
        "journal_format": "Generic Academic",
        "export_docx": True,
        "export_pdf": True,
        "export_latex": True,
        "step1_inputs": {"topic": "", "domain": "", "level": "Undergraduate",
                          "country": "", "rtype": "Quantitative"},
        "research_plan": {},
        "problem_statement": "",
        "objectives_text": "",
        "questions_text": "",
        "hypotheses_text": "",
        "lit_table": [],
        "lit_narrative": "",
        "research_gap": "",
        "methodology_inputs": {"approach": "Quantitative", "design": "Descriptive",
                                "sample_size": "", "sampling_method": "",
                                "data_collection": ""},
        "methodology_text": "",
        "questionnaire": [],
        "results_text": "",
        "conclusion_text": "",
        "references_text": "",
        "final_paper": {},
        "last_error": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def call_gemini(prompt, json_mode=False, temperature=0.5, max_output_tokens=2048):
    api_key = st.session_state.get("api_key", "").strip()
    if not api_key:
        return None, "Please enter your Gemini API key in the sidebar to continue."
    if not GENAI_AVAILABLE:
        return None, "google-generativeai package is not installed in this environment."
    try:
        genai.configure(api_key=api_key)
        config = {"temperature": temperature, "max_output_tokens": max_output_tokens}
        if json_mode:
            config["response_mime_type"] = "application/json"
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(prompt, generation_config=config)
        text = getattr(response, "text", None)
        if not text or not text.strip():
            return None, "The AI returned an empty response. Please try again."
        return text.strip(), None
    except Exception as e:
        return None, f"Gemini API error: {str(e)}"


def safe_json_parse(text):
    if not text:
        return None
    cleaned = text.strip()
    cleaned = re.sub(r"^```(json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        match = re.search(r"(\[.*\]|\{.*\})", cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1))
            except Exception:
                return None
    return None


def show_error(msg):
    st.session_state.last_error = msg
    st.error(msg)


def generate_research_plan(topic, domain, level, country, rtype):
    prompt = f"""You are an academic research assistant.
Topic: {topic}
Domain: {domain}
Research Level: {level}
Country: {country}
Research Type: {rtype}

Return ONLY valid JSON with exactly these keys:
{{"refined_topic": "...", "scope": "...", "keywords": ["...","..."], "suggested_direction": "..."}}
No markdown fences, no extra text."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=600)
    if err:
        return None, err
    data = safe_json_parse(text)
    if not data:
        return None, "Could not parse the AI response. Please try again."
    return data, None


def generate_problem_statement(mode, existing=None):
    plan = st.session_state.research_plan
    base = f"Refined Topic: {plan.get('refined_topic','')}\nScope: {plan.get('scope','')}\n"
    if mode == "generate":
        instr = "Write a concise academic problem statement (120-180 words)."
    elif mode == "regenerate":
        instr = "Write a fresh, differently-worded academic problem statement (120-180 words)."
    elif mode == "expand":
        instr = f"Expand this problem statement with more academic depth (200-260 words):\n{existing}"
    else:
        instr = f"Simplify this problem statement into clear beginner-friendly language (100-140 words):\n{existing}"
    prompt = base + instr + "\nReturn only the statement text, no heading."
    text, err = call_gemini(prompt, max_output_tokens=500)
    return text, err


def generate_objectives():
    plan = st.session_state.research_plan
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Scope: {plan.get('scope','')}
Problem Statement: {st.session_state.problem_statement[:500]}

Generate academic research content in markdown with three headed sections:
### Research Objectives
(3-5 concise bullet points)
### Research Questions
(3-5 concise bullet points)
### Hypotheses
(2-4 bullet points, or state "Not applicable for qualitative research" if appropriate)
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=700)
    return text, err


def generate_literature_table():
    plan = st.session_state.research_plan
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Domain Scope: {plan.get('scope','')}

Generate 7 illustrative academic literature review entries relevant to this topic.
Return ONLY a JSON array, each item with exactly these keys:
"author", "year", "methodology", "findings", "limitation".
Keep each field under 25 words. No markdown fences."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=1200)
    if err:
        return None, err
    data = safe_json_parse(text)
    if not isinstance(data, list):
        return None, "Could not parse literature review JSON. Please try again."
    return data, None


def generate_literature_narrative():
    table = st.session_state.lit_table
    compact = "\n".join(
        f"- {row.get('author','')} ({row.get('year','')}): {row.get('findings','')}"
        for row in table
    )
    prompt = f"""Studies:
{compact}

Write a cohesive academic narrative literature review (250-350 words) synthesizing
these studies, noting patterns, agreements and contrasts. Return only prose, no headings."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    return text, err


def generate_research_gap():
    narrative = st.session_state.lit_narrative[:900]
    prompt = f"""Literature Review Summary:
{narrative}

Identify research gaps. Return markdown with these headings:
### Academic Gap
### Practical Gap
### Future Research Opportunity
2-3 sentences under each heading."""
    text, err = call_gemini(prompt, max_output_tokens=600)
    return text, err


def generate_methodology():
    plan = st.session_state.research_plan
    m = st.session_state.methodology_inputs
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Approach: {m['approach']}
Design: {m['design']}
Sample Size: {m['sample_size']}
Sampling Method: {m['sampling_method']}
Data Collection Method: {m['data_collection']}

Write a complete academic methodology section (300-400 words) covering research design,
population and sample, sampling technique, data collection procedure, and data analysis
approach. Use markdown subheadings. Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    return text, err


def generate_questionnaire():
    plan = st.session_state.research_plan
    m = st.session_state.methodology_inputs
    prompt = f"""Topic: {plan.get('refined_topic','')}
Study Approach: {m['approach']}

Generate 12 five-point Likert scale questionnaire items (Strongly Disagree to Strongly Agree)
measuring constructs relevant to this topic. Return ONLY a JSON array of 12 short statement
strings (no numbering, statement form, no question marks). No markdown fences."""
    text, err = call_gemini(prompt, json_mode=True, max_output_tokens=700)
    if err:
        return None, err
    data = safe_json_parse(text)
    if not isinstance(data, list):
        return None, "Could not parse questionnaire JSON. Please try again."
    return data, None


def generate_results():
    plan = st.session_state.research_plan
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Methodology Summary: {st.session_state.methodology_text[:400]}

Generate markdown with these headings, 60-100 words each:
### Expected Findings
### Managerial Implications
### Academic Implications
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=700)
    return text, err


def generate_conclusion():
    plan = st.session_state.research_plan
    prompt = f"""Refined Topic: {plan.get('refined_topic','')}
Expected Results Summary: {st.session_state.results_text[:500]}

Generate markdown with these headings, 60-100 words each:
### Conclusion
### Recommendations
### Limitations
### Future Scope
Return only markdown."""
    text, err = call_gemini(prompt, max_output_tokens=800)
    return text, err


def generate_references():
    plan = st.session_state.research_plan
    style = st.session_state.citation_style
    table = st.session_state.lit_table
    authors = "; ".join(f"{r.get('author','')} ({r.get('year','')})" for r in table) or "general sources"
    prompt = f"""Topic: {plan.get('refined_topic','')}
Citation Style: {style}
Known authors referenced earlier: {authors}

Generate a reference list of 10 entries in {style} citation style relevant to this topic.
Return as a markdown numbered list, one reference per line, no extra commentary."""
    text, err = call_gemini(prompt, max_output_tokens=900)
    if err:
        return None, err
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if "[Verification Required]" not in line:
            line = line + " [Verification Required]"
        lines.append(line)
    return "\n".join(lines), None


def step_completed(n):
    checks = {
        1: bool(st.session_state.research_plan),
        2: bool(st.session_state.problem_statement.strip()),
        3: bool(st.session_state.objectives_text.strip()),
        4: bool(st.session_state.lit_table) and bool(st.session_state.lit_narrative.strip()),
        5: bool(st.session_state.research_gap.strip()),
        6: bool(st.session_state.methodology_text.strip()),
        7: bool(st.session_state.questionnaire),
        8: bool(st.session_state.results_text.strip()),
        9: bool(st.session_state.conclusion_text.strip()),
        10: bool(st.session_state.references_text.strip()),
    }
    return checks.get(n, False)


def progress_chart(completed, total=10):
    fig, ax = plt.subplots(figsize=(2.4, 2.4))
    remaining = total - completed
    colors_ = ["#1565C0", "#DCE8F7"]
    ax.pie([completed, remaining] if remaining > 0 else [completed, 0.0001],
           colors=colors_, startangle=90,
           wedgeprops=dict(width=0.35, edgecolor="white"))
    ax.text(0, 0, f"{completed}/{total}", ha="center", va="center",
            fontsize=15, fontweight="bold", color="#0D3B66")
    ax.set_aspect("equal")
    fig.patch.set_alpha(0)
    return fig


def render_sidebar():
    with st.sidebar:
        st.header("Research Settings")
        st.session_state.api_key = st.text_input(
            "Gemini API Key", value=st.session_state.api_key, type="password",
            help="Your key is used only for this session and never stored."
        )
        st.session_state.citation_style = st.selectbox(
            "Citation Style", CITATION_STYLES,
            index=CITATION_STYLES.index(st.session_state.citation_style)
        )
        st.session_state.journal_format = st.selectbox(
            "Journal / Output Format", JOURNAL_FORMATS,
            index=JOURNAL_FORMATS.index(st.session_state.journal_format)
        )
        st.session_state.word_count_target = st.slider(
            "Target Word Count", 2000, 15000, st.session_state.word_count_target, step=500
        )
        st.session_state.page_count_target = st.slider(
            "Target Page Count", 10, 60, st.session_state.page_count_target
        )

        st.markdown("---")
        st.subheader("Export Settings")
        c1, c2, c3 = st.columns(3)
        st.session_state.export_docx = c1.checkbox("DOCX", value=st.session_state.export_docx)
        st.session_state.export_pdf = c2.checkbox("PDF", value=st.session_state.export_pdf)
        st.session_state.export_latex = c3.checkbox("LaTeX", value=st.session_state.export_latex)

        st.markdown("---")
        st.subheader("Progress")
        completed = sum(1 for i in range(1, 11) if step_completed(i))
        colA, colB = st.columns([1, 1])
        with colA:
            st.pyplot(progress_chart(completed), use_container_width=True)
        with colB:
            st.metric("Sections Completed", f"{completed}/10")
            st.progress(completed / 10)

        st.markdown("---")
        st.markdown(
            '<div class="privacy-box">Your research data remains in your session '
            'and is not permanently stored.</div>',
            unsafe_allow_html=True,
        )


def section_header(title, subtitle=None):
    st.markdown(f"### {title}")
    if subtitle:
        st.caption(subtitle)


def render_step1():
    section_header("Step 1 — Research Topic", "Define your topic to generate a research plan.")
    with st.container(border=True):
        inputs = st.session_state.step1_inputs
        c1, c2 = st.columns(2)
        inputs["topic"] = c1.text_input("Research Topic", value=inputs.get("topic", ""))
        inputs["domain"] = c2.text_input("Domain", value=inputs.get("domain", ""))
        c3, c4, c5 = st.columns(3)
        inputs["level"] = c3.selectbox("Research Level", RESEARCH_LEVELS,
                                        index=RESEARCH_LEVELS.index(inputs.get("level", "Undergraduate")))
        inputs["country"] = c4.text_input("Country", value=inputs.get("country", ""))
        inputs["rtype"] = c5.selectbox("Research Type", RESEARCH_TYPES,
                                        index=RESEARCH_TYPES.index(inputs.get("rtype", "Quantitative")))
        st.session_state.step1_inputs = inputs

        if st.button("Generate Research Plan", type="primary"):
            if not inputs["topic"].strip():
                show_error("Please enter a research topic.")
            else:
                with st.spinner("Generating research plan..."):
                    data, err = generate_research_plan(
                        inputs["topic"], inputs["domain"], inputs["level"],
                        inputs["country"], inputs["rtype"]
                    )
                if err:
                    show_error(err)
                else:
                    st.session_state.research_plan = data
                    st.success("Research plan generated.")

        plan = st.session_state.research_plan
        if plan:
            st.markdown("---")
            plan["refined_topic"] = st.text_input("Refined Topic", value=plan.get("refined_topic", ""))
            plan["scope"] = st.text_area("Scope", value=plan.get("scope", ""), height=90)
            kw = plan.get("keywords", [])
            kw_text = st.text_input("Keywords (comma separated)", value=", ".join(kw))
            plan["keywords"] = [k.strip() for k in kw_text.split(",") if k.strip()]
            plan["suggested_direction"] = st.text_area(
                "Suggested Direction", value=plan.get("suggested_direction", ""), height=90
            )
            st.session_state.research_plan = plan


def render_step2():
    section_header("Step 2 — Problem Statement", "Generated from your refined topic.")
    with st.container(border=True):
        if not st.session_state.research_plan:
            st.info("Complete Step 1 first to generate a research plan.")
            return
        c1, c2, c3, c4 = st.columns(4)
        if c1.button("Generate"):
            with st.spinner("Generating..."):
                text, err = generate_problem_statement("generate")
            if err:
                show_error(err)
            else:
                st.session_state.problem_statement = text
        if c2.button("Regenerate"):
            with st.spinner("Regenerating..."):
                text, err = generate_problem_statement("regenerate")
            if err:
                show_error(err)
            else:
                st.session_state.problem_statement = text
        if c3.button("Expand"):
            with st.spinner("Expanding..."):
                text, err = generate_problem_statement("expand", st.session_state.problem_statement)
            if err:
                show_error(err)
            else:
                st.session_state.problem_statement = text
        if c4.button("Simplify"):
            with st.spinner("Simplifying..."):
                text, err = generate_problem_statement("simplify", st.session_state.problem_statement)
            if err:
                show_error(err)
            else:
                st.session_state.problem_statement = text

        st.session_state.problem_statement = st.text_area(
            "Problem Statement", value=st.session_state.problem_statement, height=180
        )


def render_step3():
    section_header("Step 3 — Objectives & Research Questions")
    with st.container(border=True):
        if not st.session_state.problem_statement.strip():
            st.info("Complete Step 2 first.")
            return
        if st.button("Generate Objectives, Questions & Hypotheses", type="primary"):
            with st.spinner("Generating..."):
                text, err = generate_objectives()
            if err:
                show_error(err)
            else:
                st.session_state.objectives_text = text
        st.session_state.objectives_text = st.text_area(
            "Objectives / Questions / Hypotheses", value=st.session_state.objectives_text, height=320
        )


def render_step4():
    section_header("Step 4 — Literature Review")
    with st.container(border=True):
        st.markdown(
            '<div class="verify-box">AI-generated references must be verified against original sources.</div>',
            unsafe_allow_html=True,
        )
        if not st.session_state.research_plan:
            st.info("Complete Step 1 first.")
            return
        if st.button("Generate Literature Review Table", type="primary"):
            with st.spinner("Generating literature table..."):
                data, err = generate_literature_table()
            if err:
                show_error(err)
            else:
                st.session_state.lit_table = data

        if st.session_state.lit_table:
            df = pd.DataFrame(st.session_state.lit_table)
            edited = st.data_editor(df, num_rows="dynamic", use_container_width=True)
            st.session_state.lit_table = edited.to_dict("records")

            if st.button("Generate Narrative Literature Review"):
                with st.spinner("Writing narrative..."):
                    text, err = generate_literature_narrative()
                if err:
                    show_error(err)
                else:
                    st.session_state.lit_narrative = text

            st.session_state.lit_narrative = st.text_area(
                "Narrative Literature Review", value=st.session_state.lit_narrative, height=220
            )


def render_step5():
    section_header("Step 5 — Research Gap")
    with st.container(border=True):
        if not st.session_state.lit_narrative.strip():
            st.info("Complete Step 4 first.")
            return
        if st.button("Generate Research Gap", type="primary"):
            with st.spinner("Generating..."):
                text, err = generate_research_gap()
            if err:
                show_error(err)
            else:
                st.session_state.research_gap = text
        st.session_state.research_gap = st.text_area(
            "Research Gap", value=st.session_state.research_gap, height=220
        )


def render_step6():
    section_header("Step 6 — Methodology Builder")
    with st.container(border=True):
        m = st.session_state.methodology_inputs
        c1, c2 = st.columns(2)
        m["approach"] = c1.selectbox("Approach", METHOD_APPROACHES,
                                      index=METHOD_APPROACHES.index(m.get("approach", "Quantitative")))
        m["design"] = c2.selectbox("Design", METHOD_DESIGNS,
                                    index=METHOD_DESIGNS.index(m.get("design", "Descriptive")))
        c3, c4, c5 = st.columns(3)
        m["sample_size"] = c3.text_input("Sample Size", value=m.get("sample_size", ""))
        m["sampling_method"] = c4.text_input("Sampling Method", value=m.get("sampling_method", ""))
        m["data_collection"] = c5.text_input("Data Collection Method", value=m.get("data_collection", ""))
        st.session_state.methodology_inputs = m

        if st.button("Generate Methodology", type="primary"):
            if not st.session_state.research_plan:
                show_error("Complete Step 1 first.")
            else:
                with st.spinner("Generating methodology..."):
                    text, err = generate_methodology()
                if err:
                    show_error(err)
                else:
                    st.session_state.methodology_text = text
        st.session_state.methodology_text = st.text_area(
            "Methodology", value=st.session_state.methodology_text, height=260
        )


def render_step7():
    section_header("Step 7 — Questionnaire Generator")
    with st.container(border=True):
        if not st.session_state.research_plan:
            st.info("Complete Step 1 first.")
            return
        if st.button("Generate Questionnaire", type="primary"):
            with st.spinner("Generating questionnaire..."):
                data, err = generate_questionnaire()
            if err:
                show_error(err)
            else:
                st.session_state.questionnaire = data

        if st.session_state.questionnaire:
            df = pd.DataFrame({"Question": st.session_state.questionnaire})
            edited = st.data_editor(df, num_rows="dynamic", use_container_width=True)
            st.session_state.questionnaire = [q for q in edited["Question"].tolist() if str(q).strip()]

            buf = export_questionnaire_docx(st.session_state.questionnaire)
            st.download_button(
                "Download Questionnaire (DOCX)", data=buf, file_name="questionnaire.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


def render_step8():
    section_header("Step 8 — Expected Results")
    with st.container(border=True):
        if not st.session_state.methodology_text.strip():
            st.info("Complete Step 6 first.")
            return
        if st.button("Generate Expected Results", type="primary"):
            with st.spinner("Generating..."):
                text, err = generate_results()
            if err:
                show_error(err)
            else:
                st.session_state.results_text = text
        st.session_state.results_text = st.text_area(
            "Expected Results", value=st.session_state.results_text, height=220
        )


def render_step9():
    section_header("Step 9 — Conclusion")
    with st.container(border=True):
        if not st.session_state.results_text.strip():
            st.info("Complete Step 8 first.")
            return
        if st.button("Generate Conclusion", type="primary"):
            with st.spinner("Generating..."):
                text, err = generate_conclusion()
            if err:
                show_error(err)
            else:
                st.session_state.conclusion_text = text
        st.session_state.conclusion_text = st.text_area(
            "Conclusion", value=st.session_state.conclusion_text, height=260
        )


def render_step10():
    section_header("Step 10 — References")
    with st.container(border=True):
        st.markdown(
            '<div class="verify-box">AI-generated references are never verified automatically. '
            'Each entry is tagged [Verification Required] and must be checked manually before submission.</div>',
            unsafe_allow_html=True,
        )
        if not st.session_state.lit_table:
            st.info("Complete Step 4 first.")
            return
        if st.button("Generate References", type="primary"):
            with st.spinner("Generating references..."):
                text, err = generate_references()
            if err:
                show_error(err)
            else:
                st.session_state.references_text = text
        st.session_state.references_text = st.text_area(
            "References", value=st.session_state.references_text, height=260
        )


def assemble_paper():
    plan = st.session_state.research_plan
    title = plan.get("refined_topic", "Untitled Research Paper")
    abstract = (
        f"This study examines {plan.get('refined_topic','the stated topic').lower()}. "
        f"{st.session_state.problem_statement[:250]} "
        f"The research adopts a {st.session_state.methodology_inputs.get('approach','')} "
        f"{st.session_state.methodology_inputs.get('design','').lower()} design. "
        f"{st.session_state.results_text[:200]}"
    ).strip()
    introduction = f"{plan.get('scope','')}\n\n{st.session_state.problem_statement}".strip()

    paper = {
        "Title": title,
        "Abstract": abstract,
        "Introduction": introduction,
        "Problem Statement": st.session_state.problem_statement,
        "Objectives, Questions & Hypotheses": st.session_state.objectives_text,
        "Literature Review": st.session_state.lit_narrative,
        "Research Gap": st.session_state.research_gap,
        "Methodology": st.session_state.methodology_text,
        "Questionnaire": "\n".join(f"{i+1}. {q}" for i, q in enumerate(st.session_state.questionnaire)),
        "Expected Results": st.session_state.results_text,
        "Conclusion": st.session_state.conclusion_text,
        "References": st.session_state.references_text,
    }
    return paper


def run_quality_checks():
    checks = []
    p = st.session_state
    for i, label in enumerate(
        ["Topic & Plan", "Problem Statement", "Objectives", "Literature Review",
         "Research Gap", "Methodology", "Questionnaire", "Expected Results",
         "Conclusion", "References"], start=1
    ):
        checks.append((f"Section: {label}", step_completed(i)))

    checks.append(("Methodology present", bool(p.methodology_text.strip())))
    checks.append(("References present", bool(p.references_text.strip())))
    checks.append(("Objectives present", bool(p.objectives_text.strip())))
    checks.append(("Problem statement length adequate (>40 words)",
                    len(p.problem_statement.split()) > 40))
    checks.append(("Conclusion length adequate (>40 words)",
                    len(p.conclusion_text.split()) > 40))
    checks.append(("Literature review has at least 3 entries", len(p.lit_table) >= 3))
    checks.append(("Questionnaire has at least 8 items", len(p.questionnaire) >= 8))
    return checks


def render_quality_check():
    section_header("Quality Check Panel")
    with st.container(border=True):
        checks = run_quality_checks()
        passed = sum(1 for _, ok in checks if ok)
        st.metric("Quality Score", f"{passed}/{len(checks)}")
        st.progress(passed / len(checks))
        for label, ok in checks:
            icon = "✅" if ok else "⚠️"
            st.markdown(f"{icon} {label}")


def render_final_paper():
    section_header("Final Research Paper", "Review and edit before exporting.")
    with st.container(border=True):
        if st.button("Generate Final Research Paper", type="primary"):
            st.session_state.final_paper = assemble_paper()
            st.success("Paper assembled below. Review and edit each section.")

        if not st.session_state.final_paper:
            st.info("Click 'Generate Final Research Paper' to assemble all completed sections.")
            return

        paper = st.session_state.final_paper
        for key in list(paper.keys()):
            with st.expander(key, expanded=False):
                paper[key] = st.text_area(key, value=paper[key], height=160, label_visibility="collapsed")
        st.session_state.final_paper = paper

        st.markdown("---")
        st.subheader("Export")
        c1, c2, c3 = st.columns(3)
        if st.session_state.export_docx:
            docx_buf = export_docx(paper, st.session_state.journal_format)
            c1.download_button(
                "Download DOCX", data=docx_buf, file_name="research_paper.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if st.session_state.export_pdf:
            pdf_buf = export_pdf(paper)
            c2.download_button("Download PDF", data=pdf_buf, file_name="research_paper.pdf",
                                mime="application/pdf")
        if st.session_state.export_latex:
            zip_buf = export_latex_zip(paper, st.session_state.lit_table)
            c3.download_button("Download LaTeX ZIP", data=zip_buf, file_name="research_paper_latex.zip",
                                mime="application/zip")


def journal_font(journal_format):
    mapping = {
        "IEEE": ("Times New Roman", 10),
        "Springer": ("Times New Roman", 11),
        "Elsevier": ("Times New Roman", 11),
        "College Project": ("Calibri", 12),
        "Generic Academic": ("Times New Roman", 12),
    }
    return mapping.get(journal_format, ("Times New Roman", 12))


def export_docx(paper, journal_format):
    font_name, font_size = journal_font(journal_format)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(paper.get("Title", "Research Paper"))
    run.bold = True
    run.font.size = Pt(font_size + 6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{journal_format} Format · Generated {datetime.now().strftime('%B %Y')}")
    sub_run.italic = True
    sub_run.font.size = Pt(font_size - 1)
    doc.add_page_break()

    for key, content in paper.items():
        if key == "Title":
            continue
        heading = doc.add_heading(key, level=1)
        for r in heading.runs:
            r.font.color.rgb = RGBColor(0x0D, 0x3B, 0x66)
        for para in str(content).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_questionnaire_docx(questions):
    doc = Document()
    doc.add_heading("Research Questionnaire", level=1)
    doc.add_paragraph("Please rate each statement on a 5-point scale: "
                       "1 = Strongly Disagree, 5 = Strongly Agree.")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q}    [1] [2] [3] [4] [5]")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_pdf(paper):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                             leftMargin=0.9 * inch, rightMargin=0.9 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"],
                                  textColor=colors.HexColor("#0D3B66"), alignment=TA_CENTER)
    heading_style = ParagraphStyle("HeadingStyle", parent=styles["Heading2"],
                                    textColor=colors.HexColor("#1565C0"), spaceBefore=14)
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"],
                                 alignment=TA_JUSTIFY, fontSize=10.5, leading=15)

    flow = [Paragraph(paper.get("Title", "Research Paper"), title_style), Spacer(1, 16)]
    for key, content in paper.items():
        if key == "Title":
            continue
        flow.append(Paragraph(key, heading_style))
        for para in str(content).split("\n"):
            if para.strip():
                safe = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                flow.append(Paragraph(safe, body_style))
                flow.append(Spacer(1, 6))
        flow.append(PageBreak())

    doc.build(flow)
    buf.seek(0)
    return buf


def bib_key(author, year):
    base = re.sub(r"[^A-Za-z]", "", str(author).split(",")[0].split(" ")[0]) or "Ref"
    return f"{base}{year}"


def export_latex_zip(paper, lit_table):
    title = paper.get("Title", "Research Paper").replace("&", "and")

    body_sections = []
    for key, content in paper.items():
        if key in ("Title", "References"):
            continue
        safe_content = str(content).replace("%", "\\%").replace("_", "\\_").replace("&", "\\&")
        body_sections.append(f"\\section{{{key}}}\n{safe_content}\n")

    cite_keys = [bib_key(r.get("author", "Author"), r.get("year", "n.d.")) for r in lit_table]
    cite_str = ", ".join(cite_keys) if cite_keys else ""

    main_tex = f"""\\documentclass[12pt]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{cite}}
\\title{{{title}}}
\\author{{}}
\\date{{}}
\\begin{{document}}
\\maketitle

{"".join(body_sections)}

\\section{{References}}
This study draws on prior work \\cite{{{cite_str}}}. See references.bib.
All entries require manual verification.

\\bibliographystyle{{apalike}}
\\bibliography{{references}}
\\end{{document}}
"""

    bib_entries = []
    for row in lit_table:
        key = bib_key(row.get("author", "Author"), row.get("year", "n.d."))
        bib_entries.append(
            f"@article{{{key},\n"
            f"  author = {{{row.get('author','Unknown')}}},\n"
            f"  year = {{{row.get('year','n.d.')}}},\n"
            f"  title = {{{row.get('findings','Untitled')[:60]}}},\n"
            f"  note = {{Verification Required}}\n}}\n"
        )
    references_bib = "\n".join(bib_entries) if bib_entries else (
        "% No literature entries available.\n% Verification Required for all sources.\n"
    )

    readme = (
        "AI Research Studio - LaTeX Export\n"
        "==================================\n\n"
        "Files included:\n"
        "  main.tex        - Main paper source\n"
        "  references.bib  - Bibliography entries (AI-generated, unverified)\n\n"
        "To compile: run pdflatex main.tex followed by bibtex main, then pdflatex main.tex twice.\n\n"
        "IMPORTANT: All references and citations were generated by AI and must be manually\n"
        "verified against original academic sources before submission.\n"
    )

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("main.tex", main_tex)
        zf.writestr("references.bib", references_bib)
        zf.writestr("README.txt", readme)
    zip_buf.seek(0)
    return zip_buf


def render_header():
    st.markdown(
        """
        <div class="app-header">
            <h1>AI Research Studio</h1>
            <p>A guided AI assistant for first-time academic research papers — from topic to final draft.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def main():
    st.set_page_config(page_title="AI Research Studio", page_icon="📚", layout="wide")
    inject_css()
    init_session_state()
    render_header()
    render_sidebar()

    tabs = st.tabs(STEP_TITLES)
    with tabs[0]:
        render_step1()
    with tabs[1]:
        render_step2()
    with tabs[2]:
        render_step3()
    with tabs[3]:
        render_step4()
    with tabs[4]:
        render_step5()
    with tabs[5]:
        render_step6()
    with tabs[6]:
        render_step7()
    with tabs[7]:
        render_step8()
    with tabs[8]:
        render_step9()
    with tabs[9]:
        render_step10()
    with tabs[10]:
        render_final_paper()
    with tabs[11]:
        render_quality_check()


if __name__ == "__main__":
    main()